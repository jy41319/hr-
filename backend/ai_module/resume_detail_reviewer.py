"""
简历逐段批注器 - 对简历各段落进行详细审查并生成批注
"""
import os
import io
import sys
import shutil
from datetime import datetime
from typing import List, Dict, Any, Tuple
from concurrent.futures import ThreadPoolExecutor, as_completed

from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

try:
    import fitz
    HAS_FITZ = True
except ImportError:
    HAS_FITZ = False

from .resume_evaluator import ResumeEvaluator
from .token_counter import TokenAccumulator, accumulate_model_tokens, extract_token_usage
from .resume_structure import get_resume_structure_extractor, save_structure_debug
from . import resume_prompts
from task_control import TaskCancelledError, can_attach_log, ensure_task_active


class ResumeDetailReviewer:
    """
    简历逐段批注器
    对简历的每个段落进行详细审查，识别问题并生成批注
    """

    def __init__(self, profile_config: Dict[str, Any] = None):
        self.ai_evaluator = ResumeEvaluator(profile_config)
        self.llm = self.ai_evaluator.llm
        self.llm_structured = self.ai_evaluator.llm_structured
        self.position_type = self.ai_evaluator.position_type
        self.structure_extractor = get_resume_structure_extractor()

    def evaluate_sections_batch(self, sections: List[Tuple[int, str, str]]) -> Tuple[List[Dict[str, Any]], int]:
        if not sections:
            return [], 0

        sections_text = ""
        for i, (idx, text, sec_type) in enumerate(sections):
            sections_text += f"\n**段落 {i} (索引 {idx}, 类型 {sec_type}):**\n{text}\n"

        prompt = f"""**背景信息：当前时间为{datetime.now().year}年。**

你是一位资深HR简历审查专家。请仔细审查以下 {len(sections)} 个简历段落，识别每个段落中存在的问题。

{sections_text}

**审查要求:**
1. 检查每个段落是否存在以下问题：
   - 时间矛盾：工作/学历时间重叠或顺序混乱
   - 夸大表述：明显夸大能力或成就（如"精通10种编程语言"）
   - 格式问题：排版混乱、标点不规范、错别字
   - 关键信息缺失：缺少联系方式、学历、关键工作细节
   - 语法错误：语句不通、用词不当
   - 逻辑问题：描述前后矛盾、因果关系不清
   - 表述问题：表达模糊、冗余啰嗦

2. **你必须在 comment 字段中给出具体评价内容，禁止留空**

3. 严重程度：critical(致命)/major(严重)/moderate(中等)/minor(轻微)/none(无问题)

请严格按JSON格式返回 {len(sections)} 个审查结果。
"""

        try:
            structured_llm = self.llm_structured.with_structured_output(resume_prompts.BatchSectionReview, include_raw=True)
            raw_resp = structured_llm.invoke(prompt)
            result = raw_resp['parsed']
            result_list = result.reviews

            output_text_est = str([r.dict() for r in result_list])
            input_tokens, output_tokens = extract_token_usage(raw_resp.get('raw'), prompt, output_text_est)
            tokens = input_tokens + output_tokens
            accumulate_model_tokens(input_tokens, output_tokens)

            result_dict = {r.section_index: r for r in result_list}

            final_results = []
            for i, (idx, text, sec_type) in enumerate(sections):
                if idx in result_dict:
                    review = result_dict[idx].dict()
                else:
                    review = {
                        'section_index': idx,
                        'has_issues': False,
                        'issue_type': '无问题',
                        'severity': 'none',
                        'comment': '该段落无明显问题',
                        'suggestion': ''
                    }

                if not review.get('comment') or review['comment'].strip() == '':
                    if review.get('has_issues'):
                        review['comment'] = f"发现{review.get('issue_type', '问题')}，需要改进"
                    else:
                        review['comment'] = "该段落无明显问题"

                final_results.append(review)
            return final_results, tokens
        except Exception as e:
            print(f"[ERROR] 批量段落审查失败: {e}")
            return [{
                'section_index': idx,
                'has_issues': False,
                'issue_type': '无问题',
                'severity': 'none',
                'comment': '审查过程出错，请人工检查',
                'suggestion': ''
            } for idx, _, _ in sections], 0

    def extract_sections(self, doc_path: str):
        structure = self.structure_extractor.extract(doc_path)
        return [(section.text, section.section_type) for section in structure.all_sections]

    def _build_comment_text(self, review: Dict[str, Any]) -> str:
        parts = [f"【{review['issue_type']}】"]
        if review.get('comment'):
            parts.append(f"\n批注: {review['comment']}")
        if review.get('suggestion'):
            parts.append(f"\n\n改进建议: {review['suggestion']}")
        return ''.join(parts)

    def _add_comment_to_paragraph(self, paragraph, comment_text: str, comment_id: int, author: str = "HR审查助手"):
        p = paragraph._element
        if len(paragraph.runs) == 0:
            paragraph.add_run()
        comment_start = parse_xml(f'<w:commentRangeStart {nsdecls("w")} w:id="{comment_id}"/>')
        p.insert(0, comment_start)
        comment_end = parse_xml(f'<w:commentRangeEnd {nsdecls("w")} w:id="{comment_id}"/>')
        p.append(comment_end)
        comment_reference = parse_xml(f'<w:r {nsdecls("w")}><w:commentReference w:id="{comment_id}"/></w:r>')
        p.append(comment_reference)

    def add_comments_to_document(self, doc_path: str, reviews: List[Dict[str, Any]], output_path: str, sections: list) -> None:
        ext = os.path.splitext(doc_path)[1].lower()
        if ext == '.pdf':
            self._add_comments_to_pdf(doc_path, reviews, output_path, sections)
        else:
            self._add_comments_to_docx(doc_path, reviews, output_path, sections)

    def _add_comments_to_docx(self, doc_path: str, reviews: List[Dict[str, Any]], output_path: str, sections: list) -> None:
        review_map = {}
        for review in reviews:
            if not review.get('has_issues'):
                continue
            section_idx = review['section_index']
            section = sections[section_idx]
            for source_ref in section.source_refs:
                if source_ref.kind == 'docx':
                    review_map[source_ref.raw_index] = review

        if not review_map:
            shutil.copy2(doc_path, output_path)
            return

        doc = Document(doc_path)
        raw_para_index = 0
        comment_id = 0
        comment_mappings = []

        severity_colors = {
            'critical': WD_COLOR_INDEX.RED,
            'major': WD_COLOR_INDEX.PINK,
            'moderate': WD_COLOR_INDEX.YELLOW,
            'minor': WD_COLOR_INDEX.TURQUOISE
        }

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                raw_para_index += 1
                continue

            if raw_para_index in review_map:
                review = review_map[raw_para_index]
                color = severity_colors.get(review['severity'], WD_COLOR_INDEX.YELLOW)

                if len(para.runs) == 0:
                    para.add_run(text)
                for run in para.runs:
                    run.font.highlight_color = color

                comment_text = self._build_comment_text(review)
                self._add_comment_to_paragraph(para, comment_text, comment_id)
                comment_mappings.append((comment_id, comment_text))
                comment_id += 1

            raw_para_index += 1

        temp_path = output_path + '.tmp.docx'
        doc.save(temp_path)
        self._add_comments_xml_to_docx(temp_path, output_path, comment_mappings)
        if os.path.exists(temp_path):
            os.remove(temp_path)

    def _add_comments_to_pdf(self, doc_path: str, reviews: List[Dict[str, Any]], output_path: str, sections: list) -> None:
        if not HAS_FITZ:
            raise ImportError("PyMuPDF未安装")

        issue_reviews = [r for r in reviews if r.get('has_issues')]
        if not issue_reviews:
            shutil.copy2(doc_path, output_path)
            return

        doc = fitz.open(doc_path)
        try:
            severity_colors = {
                'critical': (1.0, 0.4, 0.4),
                'major': (1.0, 0.6, 0.6),
                'moderate': (1.0, 0.9, 0.4),
                'minor': (0.6, 0.9, 1.0),
            }

            for review in issue_reviews:
                section = sections[review['section_index']]
                refs_by_page = {}
                for source_ref in section.source_refs:
                    if source_ref.kind != 'pdf' or source_ref.page_index is None or not source_ref.bbox:
                        continue
                    refs_by_page.setdefault(source_ref.page_index, []).append(fitz.Rect(source_ref.bbox))

                if not refs_by_page:
                    continue

                comment_text = self._build_comment_text(review)
                color = severity_colors.get(review.get('severity'), (1.0, 1.0, 0.4))

                for page_index, rects in refs_by_page.items():
                    page = doc[page_index]
                    for rect in rects:
                        annot = page.add_highlight_annot(rect)
                        annot.set_colors(stroke=color)
                        annot.update()

                    note_point = fitz.Point(rects[0].x1 + 8, rects[0].y0)
                    text_annot = page.add_text_annot(note_point, comment_text)
                    text_annot.set_info(title="HR审查助手")
                    text_annot.update()

            if os.path.exists(output_path):
                os.remove(output_path)
            doc.save(output_path)
        finally:
            doc.close()

    def _add_comments_xml_to_docx(self, source_path: str, dest_path: str, comment_mappings: List[Tuple[int, str]]) -> None:
        import re
        import zipfile

        comments_xml = '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
        comments_xml += '<w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">\n'

        for comment_id, comment_text in comment_mappings:
            comment_text_escaped = comment_text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            comments_xml += f'  <w:comment w:id="{comment_id}" w:author="HR审查助手" w:date="{datetime.now().isoformat()}" w:initials="HR">\n'
            comments_xml += '    <w:p><w:r><w:t xml:space="preserve">'
            comments_xml += f'{comment_text_escaped}</w:t></w:r></w:p>\n'
            comments_xml += '  </w:comment>\n'

        comments_xml += '</w:comments>'

        with zipfile.ZipFile(source_path, 'r') as source_zip:
            with zipfile.ZipFile(dest_path, 'w', zipfile.ZIP_DEFLATED) as dest_zip:
                for item in source_zip.infolist():
                    if item.filename not in ['word/comments.xml', 'word/_rels/document.xml.rels', '[Content_Types].xml']:
                        dest_zip.writestr(item, source_zip.read(item.filename))

                dest_zip.writestr('word/comments.xml', comments_xml.encode('utf-8'))

                rels_content = source_zip.read('word/_rels/document.xml.rels').decode('utf-8')
                if 'comments.xml' not in rels_content:
                    rids = re.findall(r'rId(\d+)', rels_content)
                    max_rid = max([int(rid) for rid in rids]) if rids else 0
                    new_rid = f"rId{max_rid + 1}"
                    comment_rel = f'<Relationship Id="{new_rid}" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments" Target="comments.xml"/>'
                    rels_content = rels_content.replace('</Relationships>', f'{comment_rel}</Relationships>')
                dest_zip.writestr('word/_rels/document.xml.rels', rels_content.encode('utf-8'))

                content_types = source_zip.read('[Content_Types].xml').decode('utf-8')
                if 'wordprocessingml.comments' not in content_types:
                    override = '<Override PartName="/word/comments.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"/>'
                    content_types = content_types.replace('</Types>', f'{override}</Types>')
                dest_zip.writestr('[Content_Types].xml', content_types.encode('utf-8'))

    def evaluate_document(self, doc_path: str, resume_id: int, db_session: Any, cancel_check=None) -> Tuple[List[Dict[str, Any]], str, int]:
        print(f"\n{'='*60}")
        print(f"开始逐段细评简历: {os.path.basename(doc_path)}")
        print(f"{'='*60}\n")

        if cancel_check: cancel_check()

        print("[步骤 1/4] 提取简历段落...")
        structure = self.structure_extractor.extract(doc_path)
        sections = structure.all_sections
        debug_dir = os.path.join(os.path.dirname(doc_path), 'debug')
        debug_path = os.path.join(debug_dir, f'resume_{resume_id}_structure_debug.json')
        save_structure_debug(structure, debug_path)
        print(f"✅ 共提取 {len(sections)} 个段落\n")

        if cancel_check: cancel_check()

        print("[步骤 2/4] 批量审查...")
        reviews = []
        sections_to_review = []
        for i, section in enumerate(sections):
            text = section.text
            sec_type = section.section_type
            if len(text.strip()) < 10:
                continue
            sections_to_review.append((i, text, sec_type))

        batch_size = 3
        total_batches = (len(sections_to_review) + batch_size - 1) // batch_size
        batches = []
        for batch_num in range(total_batches):
            start_idx = batch_num * batch_size
            end_idx = min(start_idx + batch_size, len(sections_to_review))
            batch = sections_to_review[start_idx:end_idx]
            batches.append((batch_num, batch))

        print(f"   共 {total_batches} 个批次...")
        token_acc = TokenAccumulator()

        def process_batch(batch_info):
            batch_num, batch = batch_info
            try:
                batch_results, batch_tokens = self.evaluate_sections_batch(batch)
                token_acc.add(batch_tokens)
                batch_reviews = []
                issues_in_batch = 0
                for (idx, text, sec_type), review_result in zip(batch, batch_results):
                    review_data = {
                        'section_index': idx,
                        'section_type': sec_type,
                        'section_text': text,
                        'has_issues': review_result.get('has_issues', False),
                        'issue_type': review_result.get('issue_type', '无问题'),
                        'severity': review_result.get('severity', 'none'),
                        'comment': review_result.get('comment', ''),
                        'suggestion': review_result.get('suggestion', '')
                    }
                    batch_reviews.append(review_data)
                    if review_result.get('has_issues'):
                        issues_in_batch += 1
                return (batch_num, batch_reviews, issues_in_batch, None)
            except Exception as e:
                return (batch_num, [], 0, str(e))

        executor = ThreadPoolExecutor(max_workers=min(total_batches, 10))
        try:
            future_to_batch = {executor.submit(process_batch, b): b for b in batches}
            for future in as_completed(future_to_batch):
                if cancel_check: cancel_check()
                batch_num, batch_reviews, issues, error = future.result()
                reviews.extend(batch_reviews)
                if error:
                    print(f"   ✗ 批次 {batch_num + 1} 失败: {error}")
                else:
                    print(f"   ✓ 批次 {batch_num + 1} - 发现 {issues} 处问题")
        except TaskCancelledError:
            executor.shutdown(wait=False, cancel_futures=True)
            raise
        finally:
            executor.shutdown(wait=False, cancel_futures=True)

        issues_count = sum(1 for r in reviews if r['has_issues'])
        print(f"\n✅ 审查完成，共发现 {issues_count} 处问题\n")

        if cancel_check: cancel_check()

        print("[步骤 3/4] 生成带批注文档...")
        from resume_models import Resume
        resume = db_session.get(Resume, resume_id)

        ext = os.path.splitext(doc_path)[1].lower()
        doc_ext = '.pdf' if ext == '.pdf' else '.docx'
        if resume and resume.candidate_name:
            safe_name = resume.candidate_name.replace('/', '_').replace('\\', '_')
            output_filename = f"{safe_name}-HR批注{doc_ext}"
        else:
            base = os.path.splitext(os.path.basename(doc_path))[0]
            output_filename = f"annotated_{base}{doc_ext}"

        output_dir = os.path.join(os.path.dirname(doc_path), 'annotated')
        os.makedirs(output_dir, exist_ok=True)
        output_path = os.path.join(output_dir, output_filename)

        self.add_comments_to_document(doc_path, reviews, output_path, sections)
        print(f"✅ 带批注文档已保存: {output_path}\n")

        if cancel_check: cancel_check()

        print("[步骤 4/4] 保存批注到数据库...")
        from resume_models import ResumeReview
        ResumeReview.query.filter_by(resume_id=resume_id).delete()

        saved_count = 0
        for review in reviews:
            if review['has_issues']:
                detailed_review = ResumeReview(
                    resume_id=resume_id,
                    section_index=review['section_index'],
                    section_type=review['section_type'],
                    section_text=review['section_text'],
                    comment=review['comment'],
                    issue_type=review['issue_type'],
                    severity=review['severity'],
                    suggestion=review.get('suggestion', '')
                )
                db_session.add(detailed_review)
                saved_count += 1

        if cancel_check: cancel_check()
        db_session.commit()
        print(f"✅ 已保存 {saved_count} 条批注到数据库\n")

        return reviews, output_path, token_acc.total


class _TeeStream:
    def __init__(self, *streams):
        self.streams = streams
    def write(self, data):
        for s in self.streams: s.write(data)
    def flush(self):
        for s in self.streams: s.flush()


def run_detailed_review_in_background(app, db, resume_id: int, task_token: str):
    log_buffer = io.StringIO()
    original_stdout = sys.stdout
    sys.stdout = _TeeStream(original_stdout, log_buffer)

    print(f"\n[后台任务] 开始逐段细评，简历 ID: {resume_id}")

    with app.app_context():
        from resume_models import Resume
        from .profile_resolver import resolve_profile_for_resume

        resume = db.session.get(Resume, resume_id)
        if not resume:
            sys.stdout = original_stdout
            return

        try:
            ensure_task_active(db.session, resume, 'detailed_review', task_token)
            resume.detailed_review_status = 'processing'
            db.session.commit()

            profile, profile_config = resolve_profile_for_resume(db.session, resume)
            reviewer = ResumeDetailReviewer(profile_config=profile_config)
            reviews, annotated_path, total_tokens = reviewer.evaluate_document(
                resume.resume_url, resume.id, db.session,
                cancel_check=lambda: ensure_task_active(db.session, resume, 'detailed_review', task_token)
            )

            ensure_task_active(db.session, resume, 'detailed_review', task_token)
            resume.detailed_review_status = 'completed'
            resume.annotated_document_url = annotated_path
            resume.tokens_used = (resume.tokens_used or 0) + total_tokens
            resume.detailed_review_error_message = None
            resume.detailed_review_task_token = None
            db.session.commit()
            print(f"[后台任务] ✅ 简历 ID: {resume_id} 逐段细评完成。")

        except TaskCancelledError as e:
            db.session.rollback()
            print(f"[后台任务] 简历 ID: {resume_id} 已终止: {e}")

        except Exception as e:
            from .document_reader import classify_file_error
            friendly_msg = classify_file_error(e)
            resume.detailed_review_status = 'failed'
            resume.detailed_review_error_message = friendly_msg
            resume.detailed_review_task_token = None
            db.session.commit()
            print(f"[后台任务] [FATAL] 简历 ID: {resume_id} 逐段细评失败: {e}")
            import traceback
            print(traceback.format_exc())
        finally:
            sys.stdout = original_stdout
            try:
                log_dir = os.path.dirname(resume.resume_url)
                os.makedirs(log_dir, exist_ok=True)
                log_path = os.path.join(log_dir, f"{resume_id}_detailed_review.log")
                with open(log_path, 'w', encoding='utf-8') as f:
                    f.write(log_buffer.getvalue())
                if can_attach_log(db.session, resume, 'detailed_review', task_token):
                    resume.detailed_review_log_url = log_path
            except Exception:
                pass
            db.session.commit()