"""
CVizr - Flask主应用
"""
import os
import re
import uuid
import threading
from datetime import datetime
from dotenv import load_dotenv
from flask import Flask, request, jsonify, session, send_file, send_from_directory
from flask_session import Session
from werkzeug.utils import secure_filename

from resume_models import db, User, Resume, ReviewProfile, ResumeReview, AigcDetection, RiskFlag, LLMModel, SystemSettings, Feedback

dotenv_path = os.path.join(os.path.dirname(__file__), '.env')
load_dotenv(dotenv_path=dotenv_path)

STORAGE_PATH = os.getenv('STORAGE_PATH', os.path.join(os.path.dirname(__file__), '..', 'storage'))
UPLOAD_FOLDER = os.getenv('UPLOAD_FOLDER', 'uploads')

app = Flask(__name__)
app.config['SECRET_KEY'] = os.getenv('SECRET_KEY', 'dev-secret-key-change-in-production')
app.config['SESSION_TYPE'] = 'filesystem'
app.config['SESSION_FILE_DIR'] = os.path.join(STORAGE_PATH, 'flask_session')
app.config['SESSION_PERMANENT'] = False

database_type = os.getenv('DATABASE_TYPE', 'sqlite').lower()
if database_type == 'postgresql':
    app.config['SQLALCHEMY_DATABASE_URI'] = os.getenv('POSTGRES_URI', 'postgresql://localhost/hr_resume_bot')
else:
    db_path = os.path.join(STORAGE_PATH, 'hr_resume_bot.db')
    app.config['SQLALCHEMY_DATABASE_URI'] = f'sqlite:///{db_path}'

app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['MAX_CONTENT_LENGTH'] = int(os.getenv('MAX_FILE_SIZE', 52428800))

os.makedirs(os.path.join(STORAGE_PATH, 'flask_session'), exist_ok=True)
os.makedirs(os.path.join(STORAGE_PATH, UPLOAD_FOLDER), exist_ok=True)

db.init_app(app)
Session(app)

EVALUATION_CONCURRENCY = int(os.getenv('EVALUATION_CONCURRENCY', '2'))
evaluation_semaphore = threading.BoundedSemaphore(max(1, EVALUATION_CONCURRENCY))
evaluation_task_started_at = {}


def _init_default_data():
    if not User.query.filter_by(username='admin').first():
        admin = User(username='admin', password='admin123', role='admin', real_name='管理员')
        db.session.add(admin)

    if not ReviewProfile.query.filter_by(is_default=True).first():
        from ai_module.resume_prompts import get_dimension_prompt_template, get_overall_prompt_template
        dim_template = get_dimension_prompt_template()
        overall_template = get_overall_prompt_template()

        import json
        criteria_path = os.path.join(os.path.dirname(__file__), 'ai_module', 'config', 'general_resume_criteria.json')
        with open(criteria_path, 'r', encoding='utf-8') as f:
            criteria = json.load(f)

        default_profile = ReviewProfile(
            name='通用简历评审模板',
            position_type='通用岗位',
            description='适用于所有岗位的通用简历评审模板',
            evaluation_criteria=criteria['evaluation_criteria'],
            dimension_prompt_template=dim_template.messages[0].prompt.template,
            overall_prompt_template=overall_template.messages[0].prompt.template,
            creator_id=1,
            is_active=True,
            is_default=True,
        )
        db.session.add(default_profile)

    db.session.commit()


def _ensure_runtime_columns():
    """SQLite create_all does not add columns; keep local installs forward-compatible."""
    if database_type != 'sqlite':
        return

    from sqlalchemy import inspect, text
    inspector = inspect(db.engine)
    resume_columns = {column['name'] for column in inspector.get_columns('resume')}
    column_specs = {
        'evaluation_stage': 'VARCHAR(50)',
        'evaluation_progress': 'INTEGER DEFAULT 0',
        'evaluation_status_message': 'VARCHAR(300)',
        'workflow_status': "VARCHAR(50) DEFAULT 'new'",
        'hr_note': 'TEXT',
        'job_name': 'VARCHAR(120)',
        'job_description': 'TEXT',
    }
    for column, spec in column_specs.items():
        if column not in resume_columns:
            db.session.execute(text(f'ALTER TABLE resume ADD COLUMN {column} {spec}'))
    db.session.commit()


with app.app_context():
    db.create_all()
    _ensure_runtime_columns()
    _init_default_data()


def _get_current_user():
    user_id = session.get('user_id')
    if not user_id:
        return None
    return db.session.get(User, user_id)


def _require_login():
    user = _get_current_user()
    if not user:
        return jsonify({'error': '请先登录'}), 401
    return None


def _require_admin():
    user = _get_current_user()
    if not user or user.role != 'admin':
        return jsonify({'error': '需要管理员权限'}), 403
    return None


def _build_upload_filename(original_filename, fallback_prefix='resume'):
    ext = os.path.splitext(original_filename or '')[1].lower()
    base = os.path.splitext(original_filename or '')[0]
    safe_base = secure_filename(base) or fallback_prefix
    return f"{uuid.uuid4().hex[:8]}_{safe_base}{ext}"


def _validate_uploaded_document(filepath):
    from ai_module.document_reader import get_document_reader
    validation = get_document_reader().validate_file(filepath)
    if not validation.get('valid'):
        try:
            os.remove(filepath)
        except OSError:
            pass
        return validation.get('message') or '文件无法解析，请检查文件格式'
    return None


def _first_match(text, patterns):
    for pattern in patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            return match.group(1).strip(" ｜|，,;；")
    return ''


def _infer_candidate_basic_info(resume):
    """Best-effort fallback for old reports generated before candidate_basic_info existed."""
    existing = (resume.evaluation_result or {}).get('candidate_basic_info') or {}
    if existing and any(str(value or '').strip() for value in existing.values()):
        return existing

    text = _read_resume_text_safely(resume)

    school = major = education = ''
    school_line = _first_match(text, [
        r'院校学历[：:]\s*([^\n]+)',
        r'教育背景[：:]\s*([^\n]+)',
    ])
    if school_line:
        parts = [part.strip() for part in re.split(r'[/／｜|]', school_line) if part.strip()]
        if len(parts) >= 1:
            school = parts[0]
        if len(parts) >= 2:
            major = parts[1]
        if len(parts) >= 3:
            education = parts[2]

    return {
        'age': _first_match(text, [r'年龄[：:]\s*(\d{1,2}\s*岁?)', r'(\d{1,2}\s*岁)']),
        'gender': _first_match(text, [r'性别[：:]\s*([男女])']),
        'education': education or _first_match(text, [r'学历[：:]\s*([^\n｜|/／]+)', r'(本科|硕士|博士|大专)']),
        'school': school or _first_match(text, [r'(?:学校|院校)[：:]\s*([^\n｜|/／]+)']),
        'major': major or _first_match(text, [r'专业[：:]\s*([^\n｜|/／]+)']),
        'graduation_year': _first_match(text, [r'毕业(?:时间|年份)?[：:]\s*([0-9]{4}(?:\.[0-9]{1,2})?)', r'([0-9]{4}\s*届)']),
        'city': _first_match(text, [r'城市[：:]\s*([^\n｜|，,]+)', r'(?:所在城市|期望城市)[：:]\s*([^\n｜|，,]+)']),
        'work_years': _first_match(text, [r'工作年限[：:]\s*([^\n｜|，,]+)', r'(\d+\s*年(?:工作|实习)?经验)']),
        'expected_salary': _first_match(text, [r'期望薪资[：:]\s*([^\n｜|，,]+)']),
        'phone': resume.candidate_phone or _first_match(text, [r'(1[3-9]\d[-\s]?\d{4}[-\s]?\d{4})']),
        'email': resume.candidate_email or _first_match(text, [r'([A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,})']),
    }


def _read_resume_text_safely(resume):
    try:
        from ai_module.document_reader import get_document_reader
        return get_document_reader().read(resume.resume_url)
    except Exception:
        return ''


def _fallback_dimension_evidence_from_text(text, dimension_key, dimension_payload=None):
    lines = [
        line.strip(" -•\t")
        for line in re.split(r"[\n\r]+", text or "")
        if len(line.strip()) >= 6
    ]
    lines = [
        line for line in lines
        if "模拟测试数据" not in line and "以下姓名、经历、联系方式均为虚构" not in line
    ]
    keyword_map = {
        "basic_info": r"姓名|性别|年龄|城市|手机|电话|邮箱|院校|学历|专业|毕业|薪资",
        "format": r"基本信息|个人摘要|核心技能|工作|实习|项目|教育|经历|补充说明",
        "work_logic": r"工作|实习|经历|负责|项目|20\d{2}|至今|公司|团队|部门",
        "skill_match": r"技能|能力|工具|Python|SQL|Java|前端|后端|AI|AIGC|Prompt|RAG|产品|运营|设计|数据|剪辑|小红书",
        "risk_assessment": r"风险|待确认|补充说明|毕业时间|期望薪资|未体现|不符|缺少|疑似|模拟风险",
        "overall_impression": r"个人摘要|核心技能|工作|实习|项目|负责|成果|获得|排名|优化|提升",
    }
    pattern = keyword_map.get(dimension_key, keyword_map["overall_impression"])
    matched = [line for line in lines if re.search(pattern, line, re.IGNORECASE)]
    selected = (matched or lines)[:3]
    payload = dimension_payload or {}
    summary_source = payload.get("strengths") or payload.get("weaknesses") or payload.get("feedback") or "维度判断"
    summary = re.split(r"[。；;\n]", str(summary_source).strip())[0][:80] or "维度判断"
    return [{"summary": summary, "evidence": line[:260]} for line in selected]


def _enrich_evaluation_for_report(resume):
    result = dict(resume.evaluation_result or {})
    if not result:
        return result
    text = ''
    if not result.get('candidate_basic_info'):
        result['candidate_basic_info'] = _infer_candidate_basic_info(resume)
    dimensions = result.get('dimension_evaluations')
    if isinstance(dimensions, dict):
        for key, payload in dimensions.items():
            if not isinstance(payload, dict):
                continue
            if payload.get('evidence_quotes'):
                continue
            if not text:
                text = _read_resume_text_safely(resume)
            payload['evidence_quotes'] = _fallback_dimension_evidence_from_text(text, key, payload)
    return result


def _start_evaluation_task(resume):
    resume_id = resume.id
    task_token = uuid.uuid4().hex
    evaluation_task_started_at[task_token] = datetime.utcnow()
    resume.status = 'evaluating'
    resume.evaluation_stage = 'queued'
    resume.evaluation_progress = 5
    resume.evaluation_status_message = '任务已提交，正在等待评估资源'
    resume.evaluation_error_message = None
    resume.evaluation_task_token = task_token
    db.session.commit()

    from ai_module.resume_evaluator import run_evaluation_in_background

    def _queued_evaluation_runner():
        with evaluation_semaphore:
            # Start the watchdog only after the job leaves the queue, so later
            # batch items do not time out before they actually start.
            timeout_seconds = int(os.getenv('EVALUATION_TASK_TIMEOUT', '480'))
            watchdog = threading.Timer(timeout_seconds, _mark_evaluation_timeout, args=(resume_id, task_token))
            watchdog.daemon = True
            watchdog.start()
            try:
                run_evaluation_in_background(app, db, resume_id, task_token)
            finally:
                watchdog.cancel()
                evaluation_task_started_at.pop(task_token, None)

    thread = threading.Thread(
        target=_queued_evaluation_runner,
        daemon=True,
    )
    thread.start()
    return task_token


def _mark_evaluation_timeout(resume_id, task_token):
    with app.app_context():
        resume = db.session.get(Resume, resume_id)
        if not resume:
            return
        if resume.status == 'evaluating' and resume.evaluation_task_token == task_token:
            resume.status = 'failed'
            resume.evaluation_stage = 'timeout'
            resume.evaluation_progress = 100
            resume.evaluation_status_message = '模型评估超时，已转入人工复核'
            resume.workflow_status = 'needs_review'
            resume.evaluation_error_message = '模型评估超时，已转入人工复核'
            resume.evaluation_task_token = None
            evaluation_task_started_at.pop(task_token, None)
            db.session.commit()


def _release_stale_evaluations(scope_query=None):
    """Release orphaned evaluating rows left behind by killed/restarted workers."""
    stale_seconds = int(os.getenv('EVALUATION_STALE_SECONDS', str(max(540, int(os.getenv('EVALUATION_TASK_TIMEOUT', '480')) + 60))))
    cutoff = datetime.utcnow().timestamp() - stale_seconds
    query = scope_query or Resume.query
    stale_resumes = query.filter(
        Resume.status == 'evaluating',
        Resume.evaluation_stage != 'queued',
    ).all()
    changed = False
    for resume in stale_resumes:
        stale_token = resume.evaluation_task_token
        started_at = evaluation_task_started_at.get(stale_token)
        reference_time = started_at or resume.upload_time
        if not reference_time or reference_time.timestamp() > cutoff:
            continue
        resume.status = 'failed'
        resume.workflow_status = resume.workflow_status or 'needs_review'
        resume.evaluation_stage = 'timeout'
        resume.evaluation_progress = 100
        resume.evaluation_status_message = '评估任务长时间未更新，已自动转为待重新评估'
        resume.evaluation_error_message = '评估任务长时间未更新，可能是模型请求超时或后端重启导致。'
        resume.evaluation_task_token = None
        if stale_token:
            evaluation_task_started_at.pop(stale_token, None)
        changed = True
    if changed:
        db.session.commit()


def _grade_from_score(score):
    try:
        value = float(score or 0)
    except (TypeError, ValueError):
        value = 0
    if value >= 90:
        return 'A'
    if value >= 80:
        return 'B'
    if value >= 70:
        return 'C'
    if value >= 60:
        return 'D'
    return 'E'


def _extract_decision_fields(result):
    result = result or {}
    overall = result.get('overall_evaluation') or {}
    overall_score = overall.get('overall_score')
    match_score = result.get('match_score') if result.get('match_score') is not None else overall_score
    match_grade = result.get('match_grade') or _grade_from_score(match_score)
    risk_level = result.get('risk_level') or ''
    if not risk_level:
        risk_level = 'high' if (overall_score or 0) < 60 else 'medium' if (overall_score or 0) < 75 else 'low'
    recommendation = result.get('recommendation') or ''
    if not recommendation:
        if risk_level == 'high':
            recommendation = '建议人工复核'
        elif (match_score or 0) >= 75:
            recommendation = '推荐面试'
        elif (match_score or 0) >= 60:
            recommendation = '待定'
        else:
            recommendation = '建议淘汰'
    highlights = result.get('highlights') or []
    concerns = result.get('concerns') or []
    questions = result.get('interview_questions') or []
    evidence = result.get('evidence_snippets') or []
    return {
        'overallScore': overall.get('overall_score'),
        'grade': overall.get('overall_grade'),
        'matchScore': match_score,
        'matchGrade': match_grade,
        'recommendation': recommendation,
        'riskLevel': risk_level,
        'highlights': highlights,
        'concerns': concerns,
        'interviewQuestions': questions,
        'evidenceSnippets': evidence,
    }


def _safe_export_name(value, fallback='candidate'):
    safe = secure_filename(str(value or '').strip()) or fallback
    return safe[:80]


def _risk_label(value):
    return {
        'low': '低风险',
        'medium': '中风险',
        'high': '高风险',
    }.get(value or '', value or '待确认')


def _status_label(value):
    return {
        'met': '已满足',
        'partial': '部分满足',
        'missing': '缺失',
        'unknown': '待确认',
    }.get(value or '', value or '待确认')


def _set_cell_shading(cell, fill):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:fill'), fill)


def _set_cell_text(cell, text, bold=False, color=None):
    cell.text = ''
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(text or ''))
    run.bold = bold
    if color:
        run.font.color.rgb = color
    _apply_run_font(run)


def _apply_run_font(run, east_asia='Microsoft YaHei'):
    from docx.oxml.ns import qn
    run.font.name = 'Aptos'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), east_asia)


def _add_doc_heading(document, text, level=1):
    paragraph = document.add_heading('', level=level)
    run = paragraph.add_run(text)
    _apply_run_font(run)
    return paragraph


def _add_bullets(document, items, empty_text='暂无'):
    values = [str(item).strip() for item in (items or []) if str(item).strip()]
    if not values:
        values = [empty_text]
    for item in values:
        paragraph = document.add_paragraph(style='List Bullet')
        run = paragraph.add_run(item)
        _apply_run_font(run)


def _add_kv_table(document, pairs, columns=2):
    from docx.shared import Cm, RGBColor
    visible_pairs = [(label, value if value not in [None, ''] else '待确认') for label, value in pairs]
    row_count = (len(visible_pairs) + columns - 1) // columns
    table = document.add_table(rows=row_count, cols=columns * 2)
    table.style = 'Table Grid'
    table.autofit = False
    widths = [Cm(2.3), Cm(5.0)] * columns
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = widths[index]
    for idx, (label, value) in enumerate(visible_pairs):
        row = table.rows[idx // columns]
        offset = (idx % columns) * 2
        label_cell = row.cells[offset]
        value_cell = row.cells[offset + 1]
        _set_cell_shading(label_cell, 'EAF5FF')
        _set_cell_text(label_cell, label, bold=True, color=RGBColor(51, 65, 85))
        _set_cell_text(value_cell, value, color=RGBColor(15, 23, 42))
    return table


def _build_interview_summary_docx(resume):
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt, RGBColor
    from docx.oxml.ns import qn

    result = _enrich_evaluation_for_report(resume)
    overall = result.get('overall_evaluation') or {}
    decision = _extract_decision_fields(result)
    candidate_info = result.get('candidate_basic_info') or _infer_candidate_basic_info(resume)
    requirement_matches = result.get('requirement_matches') or []
    evidence = result.get('evidence_snippets') or []
    questions = result.get('interview_questions') or decision.get('interviewQuestions') or []
    highlights = result.get('highlights') or decision.get('highlights') or []
    concerns = result.get('concerns') or decision.get('concerns') or []

    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    styles = document.styles
    normal = styles['Normal']
    normal.font.name = 'Aptos'
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    for style_name, size, color in [
        ('Title', 22, '0F172A'),
        ('Heading 1', 14, '0E7490'),
        ('Heading 2', 12, '334155'),
    ]:
        style = styles[style_name]
        style.font.name = 'Aptos'
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run('CVizr 面试摘要')
    _apply_run_font(title_run)
    title_run.bold = True
    title_run.font.size = Pt(22)
    title_run.font.color.rgb = RGBColor(15, 23, 42)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(f"{resume.candidate_name or '未命名候选人'} · {datetime.now().strftime('%Y-%m-%d')}")
    _apply_run_font(subtitle_run)
    subtitle_run.font.size = Pt(10.5)
    subtitle_run.font.color.rgb = RGBColor(100, 116, 139)

    document.add_paragraph()
    _add_doc_heading(document, '一、快速决策', level=1)
    _add_kv_table(document, [
        ('建议动作', decision.get('recommendation') or '待确认'),
        ('JD匹配分', f"{decision.get('matchScore') or '-'} / {decision.get('matchGrade') or '-'}"),
        ('综合评分', f"{overall.get('overall_score', '-')} / {overall.get('overall_grade', '-')}"),
        ('风险等级', _risk_label(decision.get('riskLevel'))),
        ('处理状态', {'new': '未处理', 'shortlisted': '待面试', 'needs_review': '待复核', 'rejected': '已淘汰', 'archived': '已入库'}.get(resume.workflow_status, resume.workflow_status or '未处理')),
        ('评估时间', resume.evaluation_time.strftime('%Y-%m-%d %H:%M') if resume.evaluation_time else '待确认'),
    ])

    reason = result.get('recommendation_reason') or overall.get('overall_feedback') or ''
    if reason:
        p = document.add_paragraph()
        run = p.add_run(f"推荐理由：{reason}")
        _apply_run_font(run)
        run.bold = True
        run.font.color.rgb = RGBColor(15, 23, 42)

    _add_doc_heading(document, '二、候选人基本信息', level=1)
    _add_kv_table(document, [
        ('姓名', resume.candidate_name or '待确认'),
        ('年龄', candidate_info.get('age') or '待确认'),
        ('性别', candidate_info.get('gender') or '待确认'),
        ('最高学历', candidate_info.get('education') or '待确认'),
        ('毕业院校', candidate_info.get('school') or '待确认'),
        ('专业', candidate_info.get('major') or '待确认'),
        ('毕业年份', candidate_info.get('graduation_year') or '待确认'),
        ('城市', candidate_info.get('city') or '待确认'),
        ('工作年限', candidate_info.get('work_years') or '待确认'),
        ('期望薪资', candidate_info.get('expected_salary') or '待确认'),
        ('电话', resume.candidate_phone or candidate_info.get('phone') or '待确认'),
        ('邮箱', resume.candidate_email or candidate_info.get('email') or '待确认'),
    ])

    _add_doc_heading(document, '三、核心亮点', level=1)
    _add_bullets(document, highlights[:3], '暂无结构化亮点，建议查看完整报告。')

    _add_doc_heading(document, '四、短板与风险', level=1)
    _add_bullets(document, concerns[:3], '暂无明显短板，建议面试中继续验证。')

    if requirement_matches:
        _add_doc_heading(document, '五、JD匹配证据', level=1)
        table = document.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        headers = ['招聘需求', '状态', '简历证据', '缺口/追问']
        for i, header in enumerate(headers):
            _set_cell_shading(table.rows[0].cells[i], 'DFF7FB')
            _set_cell_text(table.rows[0].cells[i], header, bold=True, color=RGBColor(8, 89, 105))
        for item in requirement_matches[:6]:
            row = table.add_row().cells
            _set_cell_text(row[0], item.get('requirement') or '')
            _set_cell_text(row[1], _status_label(item.get('status')))
            _set_cell_text(row[2], item.get('evidence') or '待确认')
            _set_cell_text(row[3], item.get('gap') or '面试中补充确认')

    _add_doc_heading(document, '六、建议面试问题', level=1)
    for index, question in enumerate((questions or [])[:5], 1):
        paragraph = document.add_paragraph(style='List Number')
        run = paragraph.add_run(str(question))
        _apply_run_font(run)

    if evidence:
        _add_doc_heading(document, '七、需人工复核证据', level=1)
        for item in evidence[:5]:
            if not isinstance(item, dict):
                continue
            p = document.add_paragraph()
            r1 = p.add_run(f"{item.get('risk_label') or item.get('risk_type') or '风险'}：")
            _apply_run_font(r1)
            r1.bold = True
            r1.font.color.rgb = RGBColor(180, 83, 9)
            r2 = p.add_run(f"{item.get('finding') or ''}；原文：{item.get('evidence') or '待确认'}")
            _apply_run_font(r2)

    _add_doc_heading(document, '八、HR备注', level=1)
    note = resume.hr_note or '暂无备注。'
    p = document.add_paragraph()
    r = p.add_run(note)
    _apply_run_font(r)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run('由 CVizr 自动生成 · AI 结果仅供 HR 与业务面试官辅助决策')
    _apply_run_font(footer_run)
    footer_run.font.size = Pt(8.5)
    footer_run.font.color.rgb = RGBColor(148, 163, 184)

    return document


@app.route('/api/login', methods=['POST'])
def login():
    data = request.json or {}
    username = data.get('username', '')
    password = data.get('password', '')

    user = User.query.filter_by(username=username, is_active=True).first()
    if not user or user.password != password:
        return jsonify({'error': '用户名或密码错误'}), 400

    session['user_id'] = user.id
    user.last_login = datetime.utcnow()
    db.session.commit()
    return jsonify(user.to_dict())


@app.route('/api/logout', methods=['POST'])
def logout():
    session.clear()
    return jsonify({'message': '已登出'})


@app.route('/api/current-user', methods=['GET'])
def current_user():
    user = _get_current_user()
    if not user:
        return jsonify(None)
    return jsonify(user.to_dict())


@app.route('/api/profile', methods=['PUT'])
def update_profile():
    err = _require_login()
    if err: return err
    user = _get_current_user()
    data = request.json or {}
    user.real_name = data.get('realName', user.real_name)
    user.department = data.get('department', user.department)
    db.session.commit()
    return jsonify(user.to_dict())


@app.route('/api/resumes', methods=['GET'])
def get_resumes():
    err = _require_login()
    if err: return err
    user = _get_current_user()

    page = request.args.get('page', 1, type=int)
    per_page = request.args.get('perPage', 20, type=int)
    status = request.args.get('status', '')

    query = Resume.query
    if user.role != 'admin':
        query = query.filter_by(user_id=user.id)
    _release_stale_evaluations(query)
    if status:
        query = query.filter_by(status=status)

    query = query.order_by(Resume.upload_time.desc())
    total = query.count()
    resumes = query.offset((page - 1) * per_page).limit(per_page).all()

    return jsonify({
        'resumes': [r.to_dict_light() for r in resumes],
        'total': total,
        'page': page,
        'perPage': per_page,
    })


@app.route('/api/resumes/<int:id>', methods=['GET'])
def get_resume(id):
    err = _require_login()
    if err: return err
    resume = db.session.get(Resume, id)
    if not resume:
        return jsonify({'error': '简历不存在'}), 404
    payload = resume.to_dict()
    fallback_info = _infer_candidate_basic_info(resume)
    if fallback_info:
        candidate_info = dict(payload.get('candidateBasicInfo') or {})
        for key, value in fallback_info.items():
            if value and not candidate_info.get(key):
                candidate_info[key] = value
        payload['candidateBasicInfo'] = candidate_info
        payload['candidatePhone'] = payload.get('candidatePhone') or fallback_info.get('phone') or ''
        payload['candidateEmail'] = payload.get('candidateEmail') or fallback_info.get('email') or ''
    return jsonify(payload)


@app.route('/api/resumes/<int:id>/status', methods=['GET'])
def get_resume_status(id):
    err = _require_login()
    if err: return err
    _release_stale_evaluations()
    resume = db.session.get(Resume, id)
    if not resume:
        return jsonify({'error': '简历不存在'}), 404
    return jsonify({
        'id': resume.id,
        'status': resume.status,
        'stage': resume.evaluation_stage,
        'progress': resume.evaluation_progress or 0,
        'error': resume.evaluation_error_message,
        'message': resume.evaluation_status_message or ('评估任务排队或执行中。系统会限制并发，避免触发模型限流。' if resume.status == 'evaluating' else None),
        'updatedAt': resume.evaluation_time.isoformat() if resume.evaluation_time else None,
    })


@app.route('/api/resumes/<int:id>/evaluation', methods=['GET'])
def get_resume_evaluation(id):
    err = _require_login()
    if err: return err
    resume = db.session.get(Resume, id)
    if not resume:
        return jsonify({'error': '简历不存在'}), 404
    if not resume.evaluation_result:
        return jsonify({'error': '暂无评估结果'}), 404
    return jsonify(_enrich_evaluation_for_report(resume))


@app.route('/api/resumes/<int:id>/download-report', methods=['GET'])
@app.route('/api/resumes/<int:id>/download-annotated', methods=['GET'])
def download_report(id):
    err = _require_login()
    if err: return err
    resume = db.session.get(Resume, id)
    if not resume:
        return jsonify({'error': '简历不存在'}), 404

    if request.path.endswith('/download-annotated'):
        annotated = resume.annotated_document_url
        if annotated and os.path.exists(annotated):
            return send_file(annotated, as_attachment=True, download_name=os.path.basename(annotated))
        return jsonify({'error': '暂无批注文档'}), 404

    if not resume.evaluation_result:
        return jsonify({'error': '暂无可下载报告'}), 404

    report_dir = os.path.join(STORAGE_PATH, 'exports')
    os.makedirs(report_dir, exist_ok=True)
    filename = f"CVizr_interview_brief_{resume.id}_{_safe_export_name(resume.candidate_name)}_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
    filepath = os.path.join(report_dir, filename)
    document = _build_interview_summary_docx(resume)
    document.save(filepath)
    return send_file(
        filepath,
        as_attachment=True,
        download_name=filename,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    )


@app.route('/api/upload', methods=['POST'])
def upload_resume():
    err = _require_login()
    if err: return err
    user = _get_current_user()

    if 'file' not in request.files:
        return jsonify({'error': '请上传文件'}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': '请选择文件'}), 400

    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in ['.docx', '.pdf']:
        return jsonify({'error': '仅支持 .docx 和 .pdf 文件'}), 400

    filename = _build_upload_filename(file.filename)

    upload_dir = os.path.join(STORAGE_PATH, UPLOAD_FOLDER)
    os.makedirs(upload_dir, exist_ok=True)
    filepath = os.path.join(upload_dir, filename)
    file.save(filepath)

    validation_error = _validate_uploaded_document(filepath)
    if validation_error:
        return jsonify({'error': validation_error}), 400

    candidate_name = request.form.get('candidateName', '').strip() or os.path.splitext(file.filename)[0]
    profile_id = request.form.get('profileId', None, type=int)
    job_name = request.form.get('jobName', '').strip()
    job_description = request.form.get('jobDescription', '').strip()
    auto_evaluate = request.form.get('autoEvaluate', 'false').lower() == 'true'

    resume = Resume(
        user_id=user.id,
        candidate_name=candidate_name,
        resume_url=filepath,
        profile_id=profile_id,
        job_name=job_name,
        job_description=job_description,
        workflow_status='new',
    )
    db.session.add(resume)
    db.session.commit()

    if auto_evaluate:
        _start_evaluation_task(resume)

    return jsonify(resume.to_dict())


@app.route('/api/batch-upload', methods=['POST'])
def batch_upload():
    err = _require_login()
    if err: return err
    user = _get_current_user()

    files = request.files.getlist('files')
    if not files:
        return jsonify({'error': '请上传文件'}), 400

    profile_id = request.form.get('profileId', None, type=int)
    job_name = request.form.get('jobName', '').strip()
    job_description = request.form.get('jobDescription', '').strip()
    auto_evaluate = request.form.get('autoEvaluate', 'true').lower() != 'false'
    upload_dir = os.path.join(STORAGE_PATH, UPLOAD_FOLDER)
    os.makedirs(upload_dir, exist_ok=True)

    results = []
    for file in files:
        ext = os.path.splitext(file.filename)[1].lower()
        if ext not in ['.docx', '.pdf']:
            continue

        filename = _build_upload_filename(file.filename)
        filepath = os.path.join(upload_dir, filename)
        file.save(filepath)

        validation_error = _validate_uploaded_document(filepath)
        if validation_error:
            results.append({
                'filename': file.filename,
                'status': 'failed',
                'error': validation_error,
            })
            continue

        resume = Resume(
            user_id=user.id,
            candidate_name=os.path.splitext(file.filename)[0],
            resume_url=filepath,
            profile_id=profile_id,
            job_name=job_name,
            job_description=job_description,
            workflow_status='new',
        )
        db.session.add(resume)
        results.append(resume)

    db.session.commit()
    if auto_evaluate:
        for resume in [r for r in results if isinstance(r, Resume)]:
            _start_evaluation_task(resume)
    return jsonify({
        'resumes': [r.to_dict_light() for r in results if isinstance(r, Resume)],
        'failed': [r for r in results if isinstance(r, dict)],
        'count': len([r for r in results if isinstance(r, Resume)]),
    })


@app.route('/api/evaluate/<int:id>', methods=['POST'])
def evaluate_resume(id):
    err = _require_login()
    if err: return err

    resume = db.session.get(Resume, id)
    if not resume:
        return jsonify({'error': '简历不存在'}), 404
    if resume.status == 'evaluating':
        return jsonify({'error': f'简历状态为 {resume.status}，无法重复评审'}), 400

    data = request.json or {}
    if isinstance(data, dict) and data.get('jobDescription') is not None:
        resume.job_description = (data.get('jobDescription') or '').strip()
    if isinstance(data, dict) and data.get('jobName') is not None:
        resume.job_name = (data.get('jobName') or '').strip()
    resume.status = 'pending'
    resume.evaluation_stage = 'queued'
    resume.evaluation_progress = 0
    resume.evaluation_status_message = '任务已重新提交，正在等待评估资源'
    resume.evaluation_error_message = None
    resume.evaluation_task_token = None
    resume.evaluation_result = None
    resume.ai_result = None
    db.session.commit()
    task_token = _start_evaluation_task(resume)

    return jsonify({'status': 'evaluating', 'taskToken': task_token, 'resume': resume.to_dict_light()})


@app.route('/api/resumes/<int:id>/workflow', methods=['PUT'])
def update_resume_workflow(id):
    err = _require_login()
    if err: return err
    resume = db.session.get(Resume, id)
    if not resume:
        return jsonify({'error': '简历不存在'}), 404

    data = request.json or {}
    allowed = {'new', 'shortlisted', 'needs_review', 'rejected', 'archived'}
    workflow_status = data.get('workflowStatus')
    if workflow_status is not None:
        if workflow_status not in allowed:
            return jsonify({'error': '处理状态无效'}), 400
        resume.workflow_status = workflow_status
    if data.get('hrNote') is not None:
        resume.hr_note = data.get('hrNote') or ''
    db.session.commit()
    return jsonify(resume.to_dict())


@app.route('/api/resumes/<int:id>/detailed-review', methods=['POST'])
def trigger_detailed_review(id):
    err = _require_login()
    if err: return err

    resume = db.session.get(Resume, id)
    if not resume:
        return jsonify({'error': '简历不存在'}), 404

    task_token = uuid.uuid4().hex
    resume.detailed_review_status = 'processing'
    resume.detailed_review_task_token = task_token
    db.session.commit()

    from ai_module.resume_detail_reviewer import run_detailed_review_in_background
    thread = threading.Thread(
        target=run_detailed_review_in_background,
        args=(app, db, id, task_token),
        daemon=True,
    )
    thread.start()

    return jsonify({'status': 'processing', 'taskToken': task_token})


@app.route('/api/resumes/<int:id>/aigc-detection', methods=['POST'])
def trigger_aigc_detection(id):
    err = _require_login()
    if err: return err

    resume = db.session.get(Resume, id)
    if not resume:
        return jsonify({'error': '简历不存在'}), 404

    task_token = uuid.uuid4().hex
    resume.aigc_detection_status = 'processing'
    resume.aigc_detection_task_token = task_token
    db.session.commit()

    from ai_module.aigc_detector import run_aigc_detection_in_background
    thread = threading.Thread(
        target=run_aigc_detection_in_background,
        args=(app, db, id, task_token),
        daemon=True,
    )
    thread.start()

    return jsonify({'status': 'processing', 'taskToken': task_token})


@app.route('/api/resumes/<int:id>/risk-flagging', methods=['POST'])
def trigger_risk_flagging(id):
    err = _require_login()
    if err: return err

    resume = db.session.get(Resume, id)
    if not resume:
        return jsonify({'error': '简历不存在'}), 404

    task_token = uuid.uuid4().hex
    resume.risk_flagging_status = 'processing'
    resume.risk_flagging_task_token = task_token
    db.session.commit()

    from ai_module.risk_flagger import run_risk_flagging_in_background
    thread = threading.Thread(
        target=run_risk_flagging_in_background,
        args=(app, db, id, task_token),
        daemon=True,
    )
    thread.start()

    return jsonify({'status': 'processing', 'taskToken': task_token})


@app.route('/api/resumes/<int:id>/reset-evaluation', methods=['POST'])
def reset_evaluation(id):
    err = _require_admin()
    if err: return err

    resume = db.session.get(Resume, id)
    if not resume:
        return jsonify({'error': '简历不存在'}), 404

    resume.status = 'pending'
    resume.evaluation_result = None
    resume.ai_result = None
    resume.evaluation_error_message = None
    resume.evaluation_task_token = None
    db.session.commit()

    return jsonify({'message': '已重置评审状态'})


@app.route('/api/resumes/<int:id>/risk-flags', methods=['GET'])
def get_risk_flags(id):
    err = _require_login()
    if err: return err

    flags = RiskFlag.query.filter_by(resume_id=id).all()
    return jsonify([f.to_dict() for f in flags])


@app.route('/api/resumes/<int:id>/detailed-reviews', methods=['GET'])
def get_detailed_reviews(id):
    err = _require_login()
    if err: return err

    reviews = ResumeReview.query.filter_by(resume_id=id).all()
    return jsonify([r.to_dict() for r in reviews])


@app.route('/api/resumes/<int:id>/aigc-detections', methods=['GET'])
def get_aigc_detections(id):
    err = _require_login()
    if err: return err

    detections = AigcDetection.query.filter_by(resume_id=id).all()
    return jsonify([d.to_dict() for d in detections])


@app.route('/api/chat', methods=['POST'])
def chat():
    err = _require_login()
    if err: return err
    data = request.json or {}
    message = data.get('message', '')

    if not message:
        return jsonify({'error': '请输入消息'}), 400

    try:
        from ai_module.resume_evaluator import ResumeEvaluator
        evaluator = ResumeEvaluator()
        response = evaluator.llm.invoke(
            f"你是一位HR简历审查助手。请回答以下问题：\n\n{message}"
        )
        content = getattr(response, "content", str(response))
        return jsonify({'response': content})
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/profiles', methods=['GET', 'POST'])
def manage_profiles():
    if request.method == 'GET':
        profiles = ReviewProfile.query.filter_by(is_active=True).all()
        return jsonify([p.to_dict() for p in profiles])

    err = _require_admin()
    if err: return err
    data = request.json or {}

    profile = ReviewProfile(
        name=data['name'],
        position_type=data.get('positionType', '通用岗位'),
        description=data.get('description', ''),
        evaluation_criteria=data['evaluationCriteria'],
        dimension_prompt_template=data['dimensionPromptTemplate'],
        overall_prompt_template=data['overallPromptTemplate'],
        creator_id=_get_current_user().id,
    )
    db.session.add(profile)
    db.session.commit()
    return jsonify(profile.to_dict())


@app.route('/api/profiles/<int:id>', methods=['GET', 'PUT', 'DELETE'])
def profile_detail(id):
    profile = db.session.get(ReviewProfile, id)
    if not profile:
        return jsonify({'error': '模板不存在'}), 404

    if request.method == 'GET':
        return jsonify(profile.to_dict())

    err = _require_admin()
    if err: return err

    if request.method == 'PUT':
        data = request.json or {}
        profile.name = data.get('name', profile.name)
        profile.position_type = data.get('positionType', profile.position_type)
        profile.description = data.get('description', profile.description)
        profile.evaluation_criteria = data.get('evaluationCriteria', profile.evaluation_criteria)
        profile.dimension_prompt_template = data.get('dimensionPromptTemplate', profile.dimension_prompt_template)
        profile.overall_prompt_template = data.get('overallPromptTemplate', profile.overall_prompt_template)
        db.session.commit()
        return jsonify(profile.to_dict())

    if request.method == 'DELETE':
        profile.is_active = False
        db.session.commit()
        return jsonify({'message': '已删除'})


@app.route('/api/profiles/<int:id>/duplicate', methods=['POST'])
def duplicate_profile(id):
    err = _require_admin()
    if err: return err
    profile = db.session.get(ReviewProfile, id)
    if not profile:
        return jsonify({'error': '模板不存在'}), 404

    new_profile = ReviewProfile(
        name=f"{profile.name} (副本)",
        position_type=profile.position_type,
        description=profile.description,
        evaluation_criteria=profile.evaluation_criteria,
        dimension_prompt_template=profile.dimension_prompt_template,
        overall_prompt_template=profile.overall_prompt_template,
        creator_id=_get_current_user().id,
    )
    db.session.add(new_profile)
    db.session.commit()
    return jsonify(new_profile.to_dict())


@app.route('/api/llm-models', methods=['GET', 'POST'])
def manage_llm_models():
    if request.method == 'GET':
        models = LLMModel.query.all()
        return jsonify([m.to_dict() for m in models])

    err = _require_admin()
    if err: return err
    data = request.json or {}

    model = LLMModel(
        name=data['name'],
        provider=data.get('provider', ''),
        model_name=data['modelName'],
        api_base=data['apiBase'],
        api_key=data['apiKey'],
        enable_thinking=data.get('enableThinking', False),
    )
    db.session.add(model)
    db.session.commit()
    return jsonify(model.to_dict())


@app.route('/api/llm-models/<int:id>', methods=['PUT', 'DELETE'])
def llm_model_detail(id):
    err = _require_admin()
    if err: return err
    model = db.session.get(LLMModel, id)
    if not model:
        return jsonify({'error': '模型不存在'}), 404

    if request.method == 'PUT':
        data = request.json or {}
        model.name = data.get('name', model.name)
        model.provider = data.get('provider', model.provider)
        model.model_name = data.get('modelName', model.model_name)
        model.api_base = data.get('apiBase', model.api_base)
        if data.get('apiKey'):
            model.api_key = data['apiKey']
        model.enable_thinking = data.get('enableThinking', model.enable_thinking)
        db.session.commit()
        return jsonify(model.to_dict())

    db.session.delete(model)
    db.session.commit()
    return jsonify({'message': '已删除'})


@app.route('/api/llm-models/<int:id>/activate', methods=['POST'])
def activate_llm_model(id):
    err = _require_admin()
    if err: return err

    LLMModel.query.update({'is_active': False})
    model = db.session.get(LLMModel, id)
    if not model:
        return jsonify({'error': '模型不存在'}), 404
    model.is_active = True
    db.session.commit()
    return jsonify(model.to_dict())


@app.route('/api/resumes/batch-delete', methods=['POST'])
def batch_delete_resumes():
    err = _require_admin()
    if err: return err
    ids = request.json or {}
    id_list = ids.get('ids', [])
    for resume_id in id_list:
        resume = db.session.get(Resume, resume_id)
        if resume:
            db.session.delete(resume)
    db.session.commit()
    return jsonify({'message': f'已删除 {len(id_list)} 条记录'})


@app.route('/api/resumes/export-scores', methods=['POST'])
def export_scores():
    err = _require_login()
    if err: return err

    ids = (request.json or {}).get('ids', [])
    resumes = Resume.query.filter(Resume.id.in_(ids)).all() if ids else Resume.query.filter_by(status='completed').all()

    import openpyxl
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "简历评审评分表"
    headers = ['序号', '候选人', 'JD任务名称', 'JD摘要', '评审模板', '综合评分', '综合等级', 'JD匹配分', 'JD匹配等级', '建议动作', '推荐原因', '关键缺口/淘汰原因', '风险等级', '处理状态', '核心亮点', '主要短板', 'HR备注', '评审时间', '风险数']
    ws.append(headers)

    for i, r in enumerate(resumes, 1):
        result = r.evaluation_result or {}
        overall = result.get('overall_evaluation', {})
        decision = _extract_decision_fields(result)
        ws.append([
            i, r.candidate_name or '',
            r._job_display_name(),
            r._job_summary(),
            r.profile.name if r.profile else '',
            overall.get('overall_score', ''),
            overall.get('overall_grade', ''),
            decision.get('matchScore') or '',
            decision.get('matchGrade') or '',
            decision.get('recommendation') or '',
            result.get('recommendation_reason') or '',
            '；'.join((result.get('key_gaps') or []) + (result.get('knockout_reasons') or [])),
            decision.get('riskLevel') or '',
            r.workflow_status or 'new',
            '；'.join(decision.get('highlights') or []),
            '；'.join(decision.get('concerns') or []),
            r.hr_note or '',
            r.evaluation_time.isoformat() if r.evaluation_time else '',
            r.risk_flag_count or 0,
        ])

    output_dir = os.path.join(STORAGE_PATH, 'exports')
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, f'scores_{datetime.now().strftime("%Y%m%d%H%M%S")}.xlsx')
    wb.save(output_path)

    return send_file(output_path, as_attachment=True, download_name=os.path.basename(output_path))


@app.route('/api/system-settings', methods=['GET', 'PUT'])
def system_settings():
    if request.method == 'GET':
        settings = SystemSettings.query.all()
        return jsonify([s.to_dict() for s in settings])

    err = _require_admin()
    if err: return err
    data = request.json or {}
    for key, value in data.items():
        setting = SystemSettings.query.filter_by(key=key).first()
        if setting:
            setting.value = str(value)
        else:
            setting = SystemSettings(key=key, value=str(value))
            db.session.add(setting)
    db.session.commit()
    return jsonify({'message': '已更新'})


@app.route('/api/feedback', methods=['POST', 'GET'])
def feedback():
    if request.method == 'GET':
        err = _require_admin()
        if err: return err
        feedbacks = Feedback.query.order_by(Feedback.created_at.desc()).all()
        return jsonify([f.to_dict() for f in feedbacks])

    err = _require_login()
    if err: return err
    data = request.json or {}
    fb = Feedback(user_id=_get_current_user().id, content=data.get('content', ''))
    db.session.add(fb)
    db.session.commit()
    return jsonify(fb.to_dict())


@app.route('/api/feedback/<int:id>/read', methods=['PUT'])
def mark_feedback_read(id):
    err = _require_admin()
    if err: return err
    fb = db.session.get(Feedback, id)
    if fb:
        fb.status = 'read'
        db.session.commit()
    return jsonify({'message': '已标记'})


@app.route('/api/aigc-threshold', methods=['GET', 'PUT'])
def aigc_threshold():
    from resume_models import AigcThreshold
    if request.method == 'GET':
        threshold = AigcThreshold.query.first()
        if not threshold:
            threshold = AigcThreshold()
            db.session.add(threshold)
            db.session.commit()
        return jsonify(threshold.to_dict())

    err = _require_admin()
    if err: return err
    data = request.json or {}
    threshold = AigcThreshold.query.first()
    if not threshold:
        threshold = AigcThreshold()
        db.session.add(threshold)
    threshold.high_risk_threshold = float(data.get('highRiskThreshold', 80))
    threshold.medium_risk_threshold = float(data.get('mediumRiskThreshold', 60))
    threshold.overall_alert_threshold = float(data.get('overallAlertThreshold', 70))
    db.session.commit()
    return jsonify(threshold.to_dict())


@app.after_request
def after_request(response):
    origin = request.headers.get('Origin', '')
    if origin:
        response.headers.add('Access-Control-Allow-Origin', origin)
        response.headers.add('Access-Control-Allow-Credentials', 'true')
    response.headers.add('Access-Control-Allow-Headers', 'Content-Type,Authorization')
    response.headers.add('Access-Control-Allow-Methods', 'GET,POST,PUT,DELETE,OPTIONS')
    return response


if __name__ == '__main__':
    port = int(os.getenv('PORT', 5001))
    debug = os.getenv('FLASK_DEBUG', '1').lower() in {'1', 'true', 'yes'}
    app.run(host='0.0.0.0', port=port, debug=debug, use_reloader=debug)
